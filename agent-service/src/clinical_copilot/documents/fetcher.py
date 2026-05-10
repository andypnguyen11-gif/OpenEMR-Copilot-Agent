"""Document page rendering for the demo extractor (W2-03, demo cut).

The full W2-03 fetcher also pulls the Binary via FHIR with a
system-scoped JWT; this demo cut reads from a local path. Both code
paths produce the same `RenderedPage` so the rest of the extractor
does not have to know which produced it.

Supports PDFs (rendered via pypdfium2), single-image documents (PNG
/ JPG / JPEG / BMP / WEBP), multi-page TIFFs, and *synthetic-text*
formats that have no native page representation: HL7 v2 messages,
DOCX referrals, and XLSX workbooks. The synthetic path lays out
source content as monospace text on a deterministic grid so the
citation-overlay UI has something to highlight; the paired
:mod:`synthetic_render` module documents the layout constants the
extractor uses to produce bboxes that line up with the rendered
image.

Every code path returns the same ``RenderedPage`` shape so the
extractor and the citation-overlay route never branch on format.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image

from clinical_copilot.documents.synthetic_render import render_lines

# Render at 300 DPI — high enough for OCR-quality citation checks
# downstream, low enough that a 5-page lab report stays under the
# Anthropic vision per-image size budget after JPEG compression.
RENDER_DPI: int = 300

# Anthropic's vision endpoint accepts images up to 8000 px on the long
# edge. Large scanner outputs (2400+ DPI scans) blow past that; cap
# the long edge here so a 4000 px scan downsamples to something the
# API accepts without the extractor having to know about it.
MAX_LONG_EDGE_PX: int = 2400

_IMAGE_EXTS = frozenset({".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"})

# Synthetic-text extensions handled by :mod:`synthetic_render`. These
# formats have no native page representation; the renderer lays out
# their parsed contents as monospace text on a synthesized image so
# the citation-overlay UI can show something for the click-to-source
# UX. Each format's extractor publishes bboxes that line up with the
# same layout via :func:`synthetic_render.compute_line_bbox`.
_HL7_EXTS = frozenset({".hl7"})
_DOCX_EXTS = frozenset({".docx"})
_XLSX_EXTS = frozenset({".xlsx"})


@dataclass(frozen=True, slots=True)
class RenderedPage:
    """One page of a rendered document.

    ``synthetic`` is True for pages produced by
    :mod:`synthetic_render` (HL7 / docx / xlsx) rather than rasterized
    from a native page (PDF / image / TIFF). Synthetic pages carry
    their own deterministic bbox coordinates from the extractor —
    OCR-based tightening on top of them produces *worse* results
    because the line spacing (~1.5% of page height) is below the
    cross-row anchor cap, so common tokens like ``OBX`` / ``LN`` /
    ``Cholesterol`` pull bboxes across segment boundaries. Downstream
    consumers (notably :mod:`ocr_bbox`) check this flag to skip
    OCR-tightening on synthetic pages.
    """

    page_number: int  # 1-indexed
    image: Image.Image
    width_px: int
    height_px: int
    synthetic: bool = False


def render_document(path: Path) -> list[RenderedPage]:
    """Render every page of `path` to a PIL.Image.

    PDFs go through pypdfium2 at `RENDER_DPI`. Single-image documents
    (PNG, JPG, JPEG, BMP, WEBP) are loaded directly and returned as a
    one-page list. Multi-page TIFFs (fax packets — one of the cohort-5
    file formats) are iterated via Pillow's ``n_frames`` so each page
    becomes its own ``RenderedPage``. Unknown extensions fall back to
    the image loader — most scanners produce one of the listed formats.
    """

    if not path.exists():
        raise FileNotFoundError(f"document not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _render_pdf(path)
    if suffix in {".tif", ".tiff"}:
        return _render_tiff(path)
    if suffix in _IMAGE_EXTS:
        return _render_image(path)
    if suffix in _HL7_EXTS:
        return _render_hl7(path)
    if suffix in _DOCX_EXTS:
        return _render_docx(path)
    if suffix in _XLSX_EXTS:
        return _render_xlsx(path)
    raise ValueError(
        f"Unsupported document extension {suffix!r} for {path.name}. "
        f"Supported: .pdf, {sorted(_IMAGE_EXTS)}, "
        f"{sorted(_HL7_EXTS | _DOCX_EXTS | _XLSX_EXTS)}."
    )


# Backwards-compat alias — earlier W2-01/03 code imported `render_pdf`.
render_pdf = render_document


def _render_pdf(path: Path) -> list[RenderedPage]:
    pages: list[RenderedPage] = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for index, page in enumerate(pdf):
            scale = RENDER_DPI / 72.0
            bitmap = page.render(scale=scale)
            pil_image = bitmap.to_pil().convert("RGB")
            pil_image = _cap_long_edge(pil_image)
            pages.append(
                RenderedPage(
                    page_number=index + 1,
                    image=pil_image,
                    width_px=pil_image.width,
                    height_px=pil_image.height,
                )
            )
    finally:
        pdf.close()
    return pages


def _render_image(path: Path) -> list[RenderedPage]:
    image = Image.open(path).convert("RGB")
    image = _cap_long_edge(image)
    return [
        RenderedPage(
            page_number=1,
            image=image,
            width_px=image.width,
            height_px=image.height,
        )
    ]


def _render_tiff(path: Path) -> list[RenderedPage]:
    """Iterate every frame of a (possibly multi-page) TIFF.

    Fax-packet TIFFs in the cohort-5 set are 4-5 pages of bilevel
    (mode '1') CCITT-compressed scans at 1700×2200. The conversion to
    RGB is mandatory before JPEG encoding — saving a mode-'1' image as
    JPEG raises an OSError. The long-edge cap then trims the page so
    it stays within the Anthropic vision per-image budget.
    """

    pages: list[RenderedPage] = []
    with Image.open(path) as tiff:
        frame_count = getattr(tiff, "n_frames", 1)
        for index in range(frame_count):
            tiff.seek(index)
            # ``copy()`` detaches the frame from the seek cursor so the
            # subsequent iteration does not mutate the page we just
            # appended (a real Pillow gotcha on multi-frame formats).
            frame = tiff.copy().convert("RGB")
            frame = _cap_long_edge(frame)
            pages.append(
                RenderedPage(
                    page_number=index + 1,
                    image=frame,
                    width_px=frame.width,
                    height_px=frame.height,
                )
            )
    return pages


def _render_hl7(path: Path) -> list[RenderedPage]:
    """Render an HL7 v2 message as a single monospace text page.

    Splits the source on ``\\r`` (HL7's segment terminator, with
    ``\\n`` / ``\\r\\n`` tolerated for files saved by non-HL7-aware
    editors) and lays each segment on its own line — empty source
    lines included, so the segment ``line_number`` an extractor cites
    against keeps its 1-based bijection with the rendered row.

    Synthetic renders skip the long-edge cap because the renderer
    already targets a sensible page width.
    """

    raw = path.read_text(encoding="utf-8", errors="replace")
    normalized = raw.replace("\r\n", "\r").replace("\n", "\r")
    lines = normalized.split("\r")
    # Drop trailing blank lines so a file ending with a newline doesn't
    # get a blank tail row, but preserve interior blanks (rare in HL7
    # but real — preserves line_number alignment when they occur).
    while lines and lines[-1] == "":
        lines.pop()

    image = render_lines(lines).convert("RGB")
    return [
        RenderedPage(
            page_number=1,
            image=image,
            width_px=image.width,
            height_px=image.height,
            synthetic=True,
        )
    ]


def _render_docx(path: Path) -> list[RenderedPage]:
    """Render a DOCX as a single monospace text page.

    Walks every non-empty paragraph in document order and renders one
    line per paragraph. The non-empty filter matches the referral
    extractor's paragraph indexing (which also drops empty paragraphs)
    so a citation's ``page=<paragraph_index>`` lines up with the
    rendered row at that 1-based position.

    Tables, headers/footers, and inline formatting are intentionally
    flattened away — the renderer's job is to support the citation
    overlay's "show me where this came from" UX, not to reproduce
    Word's pixel-perfect layout.
    """

    doc = DocxDocument(str(path))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    image = render_lines(lines).convert("RGB")
    return [
        RenderedPage(
            page_number=1,
            image=image,
            width_px=image.width,
            height_px=image.height,
            synthetic=True,
        )
    ]


def _render_xlsx(path: Path) -> list[RenderedPage]:
    """Render every sheet of an XLSX as one synthetic page per sheet.

    Each sheet's rendered page contains one line per non-empty row in
    the form ``A1: value | B1: value | C1: value``. Empty rows are
    skipped, which matches the workbook extractor's row iteration —
    so a citation whose ``raw_text`` is ``"Sheet!B7"`` lines up with
    the row in the rendered image whose first cell is ``A7``.

    ``data_only=True`` resolves formulas to their cached results so
    the rendered text matches what a clinician would see in Excel.
    """

    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    pages: list[RenderedPage] = []
    for sheet_index, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        lines: list[str] = []
        for row in ws.iter_rows(values_only=False):
            cells_with_value = [
                (cell.coordinate, "" if cell.value is None else str(cell.value))
                for cell in row
                if cell.value not in (None, "")
            ]
            if not cells_with_value:
                continue
            lines.append(" | ".join(f"{coord}: {val}" for coord, val in cells_with_value))
        image = render_lines(lines).convert("RGB")
        pages.append(
            RenderedPage(
                page_number=sheet_index + 1,
                image=image,
                width_px=image.width,
                height_px=image.height,
                synthetic=True,
            )
        )
    wb.close()
    return pages


def _cap_long_edge(image: Image.Image) -> Image.Image:
    long_edge = max(image.width, image.height)
    if long_edge <= MAX_LONG_EDGE_PX:
        return image
    scale = MAX_LONG_EDGE_PX / long_edge
    new_size = (int(image.width * scale), int(image.height * scale))
    return image.resize(new_size, Image.Resampling.LANCZOS)


def encode_jpeg_bytes(image: Image.Image, quality: int = 85) -> bytes:
    """Encode `image` as JPEG bytes for the Anthropic vision payload."""

    buf = io.BytesIO()
    image.save(buf, format="JPEG", quality=quality)
    return buf.getvalue()


def encode_png_bytes(image: Image.Image) -> bytes:
    """Encode `image` as PNG bytes for the citation-overlay page cache.

    PNG (rather than JPEG) avoids the compression artifacts JPEG
    introduces around text glyphs and thin lines on lab printouts;
    those artifacts make bbox alignment hard to eyeball when the
    overlay UI flips a citation rectangle on hover. The size penalty
    is acceptable for the cache use case (one-time write per page,
    served on demand).
    """

    buf = io.BytesIO()
    image.save(buf, format="PNG", optimize=True)
    return buf.getvalue()
