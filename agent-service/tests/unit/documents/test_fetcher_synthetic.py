"""Format-dispatch tests for the synthetic-text branches of
:func:`clinical_copilot.documents.fetcher.render_document`.

The render output is a deterministic monospace page; these tests
focus on the dispatch + page-count contract rather than glyph-level
fidelity, which is covered by the synthetic_render unit tests.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook

from clinical_copilot.documents.fetcher import render_document
from clinical_copilot.documents.synthetic_render import (
    SYNTH_PAGE_WIDTH_PX,
    page_metrics,
)


def _write_hl7(tmp_path: Path, segments: list[str]) -> Path:
    """Write segments joined by HL7's CR terminator. Returns the path."""

    path = tmp_path / "sample.hl7"
    path.write_text("\r".join(segments) + "\r", encoding="utf-8")
    return path


def test_render_hl7_returns_one_page_at_expected_dimensions(tmp_path: Path) -> None:
    path = _write_hl7(
        tmp_path,
        [
            "MSH|^~\\&|LAB|HOSP|EHR|EHR|20260101000000||ORU^R01^ORU_R01|MSG-1|P|2.5.1",
            "PID|1||MRN-1^^^MRN^MR||DOE^JANE||19700101|F",
            "OBR|1|ORD-1|FIL-1|58410-2^CBC panel^LN",
            "OBX|1|NM|718-7^Hemoglobin^LN||13.5|g/dL|13.5-17.5|N|||F",
        ],
    )
    pages = render_document(path)
    assert len(pages) == 1
    page = pages[0]
    metrics = page_metrics(line_count=4)
    assert page.width_px == metrics.page_width_px
    assert page.height_px == metrics.page_height_px
    assert page.page_number == 1


def test_render_hl7_handles_crlf_line_endings(tmp_path: Path) -> None:
    """Files saved by non-HL7-aware editors land on disk with CRLF
    or LF instead of bare CR. The renderer must tolerate both so the
    bbox alignment with :class:`Segment.line_number` survives a
    text-editor round-trip."""

    path = tmp_path / "crlf.hl7"
    path.write_text(
        "MSH|...|MSG\r\nPID|1||MRN\r\nOBX|1|NM|||5\r\n",
        encoding="utf-8",
    )
    pages = render_document(path)
    metrics = page_metrics(line_count=3)
    assert pages[0].height_px == metrics.page_height_px


def test_render_docx_returns_one_page_per_document(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Sandia Foothills Internal Medicine")
    doc.add_paragraph("")  # empty — should be filtered out by the renderer
    doc.add_paragraph("RE: Whitaker, James | DOB: 11/22/1958 | MRN: NMM-001")
    doc.add_paragraph("Reason for Referral: hematuria workup")
    doc.save(str(path))

    pages = render_document(path)
    assert len(pages) == 1
    # Three non-empty paragraphs means three rendered lines.
    metrics = page_metrics(line_count=3)
    assert pages[0].height_px == metrics.page_height_px


def test_render_docx_raises_for_empty_document(tmp_path: Path) -> None:
    """python-docx's :class:`Document` always carries at least one
    empty paragraph, but a synthetic doc with no text content should
    still produce a renderable (mostly-blank) page rather than
    raising — the citation overlay treats empty pages as the
    legitimate "this format has no overlay-worthy content" case."""

    path = tmp_path / "empty.docx"
    Document().save(str(path))
    pages = render_document(path)
    assert len(pages) == 1
    assert pages[0].width_px == SYNTH_PAGE_WIDTH_PX


def test_render_xlsx_returns_one_page_per_sheet(tmp_path: Path) -> None:
    path = tmp_path / "sample.xlsx"
    wb = Workbook()
    sheet1 = wb.active
    assert sheet1 is not None
    sheet1.title = "Patient"
    sheet1["A1"] = "Name"
    sheet1["B1"] = "Whitaker, James"
    sheet1["A2"] = "DOB"
    sheet1["B2"] = "1958-11-22"

    sheet2 = wb.create_sheet("Medications")
    sheet2["A1"] = "Drug"
    sheet2["A2"] = "Apixaban"
    sheet2["A3"] = "Tamsulosin"
    wb.save(str(path))

    pages = render_document(path)
    assert len(pages) == 2
    assert [p.page_number for p in pages] == [1, 2]
    # Each page is sized for its own kept-row count.
    sheet1_metrics = page_metrics(line_count=2)
    sheet2_metrics = page_metrics(line_count=3)
    assert pages[0].height_px == sheet1_metrics.page_height_px
    assert pages[1].height_px == sheet2_metrics.page_height_px


def test_render_xlsx_skips_empty_rows(tmp_path: Path) -> None:
    """The renderer must skip programmatically-empty rows so its
    line index lines up with the workbook extractor's
    ``_kept_rows()`` walk. A misalignment here misplaces every
    citation on the sheet."""

    path = tmp_path / "gappy.xlsx"
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws["A1"] = "row 1"
    # Row 2 deliberately empty.
    ws["A3"] = "row 3"
    wb.save(str(path))

    pages = render_document(path)
    expected_metrics = page_metrics(line_count=2)
    assert pages[0].height_px == expected_metrics.page_height_px


def test_render_document_rejects_unknown_extension(tmp_path: Path) -> None:
    path = tmp_path / "unknown.xyz"
    path.write_text("nothing renderable here", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported document extension"):
        render_document(path)
